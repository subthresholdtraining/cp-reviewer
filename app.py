"""
Client Practical Reviewer - Streamlit App
Transforms raw grading notes into polished feedback documents.
"""

import re
import streamlit as st
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from datetime import datetime

# Page config
st.set_page_config(
    page_title="Client Practical Reviewer",
    page_icon="📝",
    layout="wide"
)

# Colors matching the SA Pro style
PURPLE = RGBColor(0x77, 0x56, 0xA7)  # #7756A7 - Purple for body text
RED = RGBColor(0xE7, 0x4E, 0x4E)  # #E74E4E - Red for headings

# Font settings
FONT_NAME = 'Lato'
TITLE_SIZE = Pt(18)
HEADING_SIZE = Pt(11)
BODY_SIZE = Pt(11)

# Words to avoid
WORDS_TO_AVOID = """
resonate, captivating, immersive, inspiring, passionate, meaningful, impactful, heartfelt,
authentic, genuine, magical, transcendent, amazing, enlightening, extraordinary, phenomenal,
profound, disrupt, blockchain, AI-powered, machine learning, neural, algorithm, interface,
user-centric, actionable, data-driven, cloud-native, agile, state-of-the-art, innovative,
groundbreaking, game-changing, disruptive, paradigm shifting, bandwidth, deliverables,
value proposition, stakeholders, alignment, strategic, mission-critical, core competency,
visibility, incentivize, ownership, value-add, leverage, utilize, synergy, blueprint, boost,
transform, transformative, bespoke, curated, premium, exclusive, elite, lifestyle, mindshare,
trending, viral, organic, artisanal, handcrafted, iconic, indispensable, unique, unprecedented,
flourish, skyrocket, methodology, framework, paradigm, conceptualize, contextualize, synthesize,
operationalize, fundamentally, inherently, ultimately, essentially, delve, discover, explore,
craft, deep dive, dive in, disruption, unicorn, moonshot, growth hacking, bleeding edge, pivot,
MVP, scale, ideate, iteration, lean, empower, unleash, unlock, Moving forward, That being said,
In terms of, With that in mind, To be honest, The fact of the matter is, At this point in time,
When all is said and done, The reality is, In my experience, instinct, gut, shedding, evoke,
enthusiast, aim, literally
"""

SYSTEM_PROMPT = """You are helping write feedback for SA Pro Trainer certification practical assessments.

Transform the raw notes into a polished review document. The tone should be:
- Conversational, like one colleague talking to another over coffee
- Warm and supportive, but not over-the-top
- Specific with examples ("Like when you...", "I noticed that...")
- Professional but friendly
- Down-to-earth and relatable

Structure the output as:
1. **What you did well** - bullet points highlighting strengths and successes
2. **What you could do differently next time** - constructive, actionable suggestions framed positively
3. **Overall** - A brief summary paragraph (2-3 sentences) capturing the key points

IMPORTANT - Preserve detail and terminology:
- Do NOT over-summarize. Each distinct point in the notes should appear in the output.
- Keep specific training terminology exactly as written: "Door is a Bore", "push-drop", "FOMO", "hyper-attachment", threshold terms, etc.
- If the notes mention a specific timestamp (e.g., "at 25 min mark"), include it
- If the notes have 10 points, the output should have roughly 10 bullet points across the sections
- Polish the language, but don't condense multiple observations into one

Style guidelines:
- Use "I noticed...", "I liked that...", "Well done for...", "Nice job..."
- Use specific observations from the notes
- Vary sentence length for natural rhythm
- Keep technical language minimal
- Use exclamation points sparingly but naturally
- End with encouragement like "Well done!" if appropriate

IMPORTANT: Avoid these words and phrases completely:
""" + WORDS_TO_AVOID + """

Here are examples of the style to match:

GOOD EXAMPLE 1:
"Great start by letting Sharon give you some background on Maggie, you used the form to help build trust and also asked some great questions."

GOOD EXAMPLE 2:
"I can see why you missed the little whine at the 25 min mark. Even if you had heard it, I think it was okay to continue given that there wasn't much left on the clock. And she didn't escalate."

GOOD EXAMPLE 3:
"Your words - 'That 2.06 seconds is less important than keeping him below threshold' are perfect!"

BAD EXAMPLE (too corporate/AI):
"Overall, this was a really strong session! Your calm, knowledgeable approach helped Desiree and John feel supported while giving them practical tools to work with."

IMPROVED VERSION:
"Great work with Desiree and John! You struck such a nice balance - professional but approachable, keeping the session relaxed while giving them really practical tools to work with."
"""


def get_polished_feedback(client, raw_notes: str, student_name: str, client_name: str, dog_name: str) -> str:
    """Use Claude to transform raw notes into polished feedback (always in English)."""

    context = f"Student being assessed: {student_name}\n"
    if client_name:
        context += f"Client in the video: {client_name}\n"
    if dog_name:
        context += f"Dog's name: {dog_name}\n"

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=2000,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": f"{context}\nHere are the raw notes from watching the assessment video:\n\n{raw_notes}\n\nPlease transform these into a polished feedback document."
            }
        ]
    )

    return message.content[0].text


def _ensure_section_headers(text: str) -> str:
    """Safety net: if Claude translated the section headers, replace them with the English originals.

    Looks for common French/Dutch translations of the three required headers and
    swaps them back to English so the Word doc parser can find them.
    """
    # Patterns that match translated versions of each header (case-insensitive)
    header_fixes = [
        {
            'english': '**What you did well**',
            'patterns': [
                r'\*\*Ce que (?:vous avez|tu as) bien fait\*\*',
                r'\*\*Wat (?:je|u) goed (?:hebt |heeft )?gedaan\*\*',
                r'\*\*(?:Points? forts?|Points? positifs?)\*\*',
                r'\*\*Wat ging er goed\*\*',
            ]
        },
        {
            'english': '**What you could do differently next time**',
            'patterns': [
                r'\*\*Ce que (?:vous pourriez|tu pourrais) faire diff[ée]remment.+?\*\*',
                r'\*\*Wat (?:je|u) de volgende keer anders (?:zou(?:dt)? )?kunnen doen\*\*',
                r'\*\*(?:Points? [àa] am[ée]liorer|Axes? d.am[ée]lioration).*?\*\*',
                r'\*\*Wat (?:je|u) anders (?:zou(?:dt)? )?kunnen doen\*\*',
            ]
        },
        {
            'english': '**Overall**',
            'patterns': [
                r'\*\*(?:En r[ée]sum[ée]|Conclusion|Bilan|Dans l.ensemble|Globalement)\*\*',
                r'\*\*(?:Algeheel|Over het geheel|Samenvatting|Algemeen|Totaal)\*\*',
            ]
        },
    ]

    for fix in header_fixes:
        # Check if the English header is already present
        if fix['english'] in text:
            continue
        # Try each translated pattern
        for pattern in fix['patterns']:
            text, count = re.subn(pattern, fix['english'], text, count=1, flags=re.IGNORECASE)
            if count > 0:
                break

    return text


def translate_feedback(client, english_text: str, target_language: str) -> str:
    """Translate the polished English feedback to French or Dutch."""

    if target_language == "French":
        language_instruction = """You are translating a dog training assessment feedback document from English to French.

CRITICAL - DO NOT TRANSLATE THESE SECTION HEADERS. Keep them EXACTLY as-is in English:
- **What you did well**
- **What you could do differently next time**
- **Overall**
Only translate the content underneath each header, not the headers themselves.

FRENCH DICTIONARY - Use these specific terms:
- Behavior consultant = consultant(e) en comportement canin
- Separation anxiety = anxiété de séparation
- Dog trainer = Consultant(e) en comportement canin
- Dog training = Éducation canine

TRANSLATION RULES:
1. Do NOT literally translate idioms - use the French equivalent instead
2. Keep these English expressions as-is: "door is a bore", "FOMO", "hyper-attachement", "push-drop"
3. Use modern, natural French - avoid antiquated expressions
4. Maintain the warm, conversational, colleague-to-colleague tone
5. Keep the educational context of dog training
6. Preserve the exact structure (bullet points, sections) of the original

Translate the following feedback to French:"""

    elif target_language == "Dutch":
        language_instruction = """You are translating a dog training assessment feedback document from English to Dutch.

CRITICAL - DO NOT TRANSLATE THESE SECTION HEADERS. Keep them EXACTLY as-is in English:
- **What you did well**
- **What you could do differently next time**
- **Overall**
Only translate the content underneath each header, not the headers themselves.

DUTCH TRANSLATION RULES:
1. Keep these English expressions as-is: "door is a bore", "FOMO", "push-drop"
2. Use "je/jij" (informal) not "u" (formal) - keep it warm and collegial
3. Avoid literal translations of English idioms - use natural Dutch equivalents
4. Watch word order in subordinate clauses (verb goes to end)
5. Use natural Dutch compound words where appropriate
6. Avoid anglicisms where good Dutch alternatives exist (e.g., use "terugkoppeling" not "feedback" if it fits naturally)
7. Keep the warm, conversational, colleague-to-colleague tone
8. Use modern Dutch - avoid stiff or formal phrasing
9. "Separation anxiety" = "verlatingsangst" or "scheidingsangst"
10. Be careful with false friends (e.g., "eventually" ≠ "eventueel")
11. Preserve the exact structure (bullet points, sections) of the original

Translate the following feedback to Dutch:"""
    else:
        return english_text

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=2000,
        messages=[
            {
                "role": "user",
                "content": f"{language_instruction}\n\n{english_text}"
            }
        ]
    )

    translated = message.content[0].text

    # Safety net: ensure all three section headers survived translation.
    # If Claude translated a header despite instructions, re-insert the English version.
    translated = _ensure_section_headers(translated)

    return translated


def create_review_document(
    student_name: str,
    review_date: str,
    reviewer_name: str,
    status: str,
    feedback_text: str
) -> BytesIO:
    """Create a Word document with the review."""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.font.color.rgb = PURPLE

    # Header - SA PRO TRAINER ASSESSMENT
    header = doc.add_paragraph()
    header_run = header.add_run("SA PRO TRAINER ASSESSMENT")
    header_run.bold = True
    header_run.font.name = FONT_NAME
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RED
    header.alignment = 1  # Center

    # Title - CLIENT PRACTICAL REVIEW (Red, 18pt, Lato)
    title = doc.add_paragraph()
    title_run = title.add_run("CLIENT PRACTICAL REVIEW")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = TITLE_SIZE
    title_run.font.color.rgb = RED
    title.alignment = 1  # Center

    doc.add_paragraph()

    # Meta info (labels in Red 11pt, values in Purple 11pt)
    meta_style = [
        ("Name: ", student_name),
        ("Date: ", review_date),
        ("Reviewer: ", reviewer_name),
        ("Status: ", status),
    ]

    for label, value in meta_style:
        p = doc.add_paragraph()
        label_run = p.add_run(label)
        label_run.bold = True
        label_run.font.name = FONT_NAME
        label_run.font.size = HEADING_SIZE
        label_run.font.color.rgb = RED
        value_run = p.add_run(value)
        value_run.font.name = FONT_NAME
        value_run.font.size = BODY_SIZE
        value_run.font.color.rgb = PURPLE

    doc.add_paragraph()

    # Parse and add the feedback sections
    lines = feedback_text.split('\n')
    current_section = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check for section headers (Red 11pt Lato)
        if '**What you did well**' in line or 'What you did well' in line:
            p = doc.add_paragraph()
            run = p.add_run("What you did well")
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_SIZE
            run.font.color.rgb = RED
            current_section = 'well'
        elif '**What you could do differently' in line or 'What you could do differently' in line:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("What you could do differently next time")
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_SIZE
            run.font.color.rgb = RED
            current_section = 'improve'
        elif '**Overall**' in line or line.startswith('Overall'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Overall")
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = HEADING_SIZE
            run.font.color.rgb = RED
            current_section = 'overall'
            # Check if there's content after "Overall" on the same line
            if ':' in line:
                content = line.split(':', 1)[1].strip()
                if content:
                    p2 = doc.add_paragraph()
                    content_run = p2.add_run(content)
                    content_run.font.name = FONT_NAME
                    content_run.font.size = BODY_SIZE
                    content_run.font.color.rgb = PURPLE
        elif line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
            # Bullet point (Purple 11pt Lato)
            content = line[2:].strip()
            # Remove markdown bold
            content = content.replace('**', '')
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            run.font.color.rgb = PURPLE
        elif current_section == 'overall' and not line.startswith('#'):
            # Overall summary paragraph (Purple 11pt Lato)
            clean_line = line.replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.font.name = FONT_NAME
            run.font.size = BODY_SIZE
            run.font.color.rgb = PURPLE

    # Save to BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer


# App UI
st.title("📝 Client Practical Reviewer")
st.markdown("*SA Pro Trainer Assessment*")
st.markdown("---")

# Get API key from secrets (no need for users to enter it)
api_key = st.secrets.get("ANTHROPIC_API_KEY", "")

# Sidebar for settings
with st.sidebar:
    st.header("Settings")
    reviewer_name = st.text_input("Reviewer Name", value="")
    st.markdown("---")
    st.markdown("### How to use")
    st.markdown("""
    1. Fill in student details
    2. Paste your raw notes
    3. Click 'Generate Review'
    4. Edit if needed
    5. Download the Word doc
    """)

# Main form
st.header("Assessment Details")

col1, col2 = st.columns(2)
with col1:
    student_name = st.text_input("Student Name", placeholder="e.g., Amanda Dwyer")
    client_name = st.text_input("Client Name (from video)", placeholder="e.g., Natalie")
with col2:
    review_date = st.date_input("Review Date", value=datetime.now())
    dog_name = st.text_input("Dog Name", placeholder="e.g., Teddy")

status = st.selectbox("Status", ["Passed", "Cleared", "Resubmit"])

st.markdown("---")
st.header("Raw Notes")
st.markdown("*Paste your notes from watching the video:*")

raw_notes = st.text_area(
    "Raw Notes",
    height=300,
    placeholder="""Nice work reassuring the client at the start.
Good coaching of the training session itself.
Great explanation of why we use Door is a Bore.
Some nervousness throughout but relaxed into it.
Ending felt a bit rushed but pacing was okay overall.
...""",
    label_visibility="collapsed"
)

st.markdown("---")

# Generate button
if st.button("✨ Generate Review", type="primary", use_container_width=True):
    if not api_key:
        st.error("API key not configured. Please contact your administrator.")
    elif not student_name:
        st.error("Please enter the student's name.")
    elif not raw_notes:
        st.error("Please enter your raw notes.")
    else:
        with st.spinner("Generating polished review..."):
            try:
                client = Anthropic(api_key=api_key)
                polished = get_polished_feedback(client, raw_notes, student_name, client_name, dog_name)
                st.session_state['polished_feedback'] = polished
                st.session_state['student_name'] = student_name
                st.session_state['review_date'] = review_date.strftime("%B %d, %Y")
                st.session_state['reviewer_name'] = reviewer_name
                st.session_state['status'] = status
                st.success("Review generated!")
            except Exception as e:
                st.error(f"Error: {e}")

# Show and edit generated feedback
if 'polished_feedback' in st.session_state:
    st.markdown("---")
    st.header("Generated Review")
    st.markdown("*Edit as needed:*")

    edited_feedback = st.text_area(
        "Polished Feedback",
        value=st.session_state['polished_feedback'],
        height=400,
        label_visibility="collapsed"
    )

    # Update session state if edited
    st.session_state['polished_feedback'] = edited_feedback

    # Translation section
    st.markdown("---")
    st.header("🌍 Translate")
    st.markdown("*Happy with the English version? Translate it:*")

    col_lang, col_btn = st.columns([1, 2])
    with col_lang:
        target_language = st.selectbox("Language", ["French", "Dutch"], label_visibility="collapsed")
    with col_btn:
        translate_clicked = st.button(f"🔄 Translate to {target_language}", use_container_width=True)

    if translate_clicked:
        if not api_key:
            st.error("API key not configured. Please contact your administrator.")
        else:
            with st.spinner(f"Translating to {target_language}..."):
                try:
                    client = Anthropic(api_key=api_key)
                    translated = translate_feedback(client, edited_feedback, target_language)
                    st.session_state['polished_feedback'] = translated
                    st.success(f"Translated to {target_language}!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Translation error: {e}")

    st.markdown("---")
    st.header("📄 Download")

    # Generate document with error handling
    try:
        doc_buffer = create_review_document(
            student_name=st.session_state['student_name'],
            review_date=st.session_state['review_date'],
            reviewer_name=st.session_state['reviewer_name'],
            status=st.session_state['status'],
            feedback_text=edited_feedback
        )

        # Get the bytes from the buffer
        doc_bytes = doc_buffer.getvalue()

        # Download button
        filename = f"Client_Practical_{st.session_state['student_name'].replace(' ', '_')}.docx"
        st.download_button(
            label="📥 Download Word Document",
            data=doc_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"Error generating document: {e}")

# Footer
st.markdown("---")
st.markdown("*Client Practical Reviewer v1.1 - SA Pro Trainer Certification*")
